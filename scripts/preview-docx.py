import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document

d = Document("section-3-evolution.docx")
n_para = sum(1 for p in d.paragraphs if p.text.strip())
n_img = 0
for shape in d.inline_shapes:
    n_img += 1
print(f"paragraphs: {n_para}, inline_images: {n_img}")

print("\n--- headings (first 20) ---")
for p in d.paragraphs[:200]:
    t = p.text.strip()
    if not t:
        continue
    # Output only short lines which are likely headings
    if len(t) < 90:
        print(t)
